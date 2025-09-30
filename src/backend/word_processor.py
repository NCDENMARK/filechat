import os
import hashlib
from pathlib import Path
from typing import List, Dict
from docx import Document

class WordProcessor:
    def __init__(self):
        # Keep track of files we've already processed
        self.processed_files = set()
    
    def scan_directory(self, directory_path: str) -> List[str]:
        """Find all Word files in a given folder"""
        word_files = []  # List to store found Word document paths
        path = Path(directory_path)  # Convert string to Path object
        
        # Check if the directory exists and is actually a directory
        if path.exists() and path.is_dir():
            # Find all .docx files (case-insensitive)
            for word_path in path.glob("*.[dD][oO][cC][xX]"):
                word_files.append(str(word_path))  # Add full path to list
            # Also check for older .doc files
            for word_path in path.glob("*.[dD][oO][cC]"):
                # Only add if it's not already in the list (avoid duplicates)
                path_str = str(word_path)
                if not path_str.endswith('.docx') and not path_str.endswith('.DOCX'):
                    word_files.append(path_str)
        
        return word_files  # Return list of Word file paths
    
    def extract_text(self, word_path: str) -> Dict:
        """Pull all text out of a Word document"""
        try:
            # Open the Word document
            doc = Document(word_path)
            text = ""  # String to accumulate all text
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += f"{paragraph.text}\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += f"{cell.text}\n"
            
            # Create a unique ID for this file using MD5 hash
            file_id = hashlib.md5(word_path.encode()).hexdigest()
            
            # Count sections/paragraphs
            section_count = len(doc.sections)
            paragraph_count = len(doc.paragraphs)
            
            # Return all the extracted information
            return {
                "file_id": file_id,  # Unique identifier
                "file_path": word_path,  # Full path to file
                "file_name": Path(word_path).name,  # Just the filename
                "text": text,  # All extracted text
                "section_count": section_count,  # Number of sections
                "paragraph_count": paragraph_count,  # Number of paragraphs
                "status": "success"  # Extraction worked
            }
        except Exception as e:
            # If something goes wrong, return error info
            return {
                "file_path": word_path,
                "file_name": Path(word_path).name,
                "status": "error",
                "error": str(e)  # What went wrong
            }
    
    def split_into_chunks(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Break long text into smaller, overlapping pieces"""
        words = text.split()  # Split text into individual words
        chunks = []  # List to store text chunks
        
        # Move through the text in steps of (chunk_size - overlap)
        for i in range(0, len(words), chunk_size - overlap):
            # Take chunk_size words starting from position i
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
        
        return chunks